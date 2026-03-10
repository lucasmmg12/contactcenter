"""
Document Processor — Extracts text from various file formats
"""
import os
import re
import csv
import io


def extract_text(file_path: str) -> str:
    """Extract text from a file based on its extension."""
    ext = os.path.splitext(file_path)[1].lower()

    extractors = {
        '.pdf': _extract_pdf,
        '.docx': _extract_docx,
        '.xlsx': _extract_excel,
        '.xls': _extract_excel,
        '.csv': _extract_csv,
    }

    extractor = extractors.get(ext, _extract_text_file)
    raw_text = extractor(file_path)
    return clean_text(raw_text)


def _extract_pdf(file_path: str) -> str:
    """Extract text from PDF using PyPDF2."""
    from PyPDF2 import PdfReader

    reader = PdfReader(file_path)
    pages = []
    for i, page in enumerate(reader.pages, 1):
        text = page.extract_text()
        if text and text.strip():
            pages.append(f"[Página {i}]\n{text.strip()}")
    return "\n\n".join(pages)


def _extract_docx(file_path: str) -> str:
    """Extract text from Word documents."""
    from docx import Document

    doc = Document(file_path)
    parts = []

    # Paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())

    # Tables
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n\n".join(parts)


def _extract_excel(file_path: str) -> str:
    """Extract text from Excel files."""
    from openpyxl import load_workbook

    wb = load_workbook(file_path, data_only=True)
    parts = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        parts.append(f"[Hoja: {sheet_name}]")
        for row in ws.iter_rows(values_only=True):
            cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def _extract_csv(file_path: str) -> str:
    """Extract text from CSV files."""
    parts = []
    with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
        reader = csv.reader(f)
        for row in reader:
            cells = [c.strip() for c in row if c.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _extract_text_file(file_path: str) -> str:
    """Extract text from plain text files (.txt, .md, .json, .xml, .html)."""
    with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
        return f.read()


def clean_text(text: str) -> str:
    """Sanitize text: remove invalid Unicode, null bytes, control characters."""
    if not text:
        return ""

    # Remove invalid Unicode surrogates
    text = text.encode('utf-8', errors='ignore').decode('utf-8', errors='ignore')

    # Remove null bytes
    text = text.replace('\x00', '')

    # Remove control characters but preserve \n and \t
    text = re.sub(r'[\x01-\x08\x0b\x0c\x0e-\x1f\x7f]', '', text)

    return text.strip()
